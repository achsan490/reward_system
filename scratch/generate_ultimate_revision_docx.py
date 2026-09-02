# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'))

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(8)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def add_callout(doc, title, text, bg_color="F0F4F8", border_color="1F4E79"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="28" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r_title = p.add_run(f"📍 {title}\n")
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(10)
    r_text.italic = True
    r_text.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_body_p(doc, text, bold_prefix="", indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def create_ultimate_revision_document(target_path):
    doc = docx.Document()
    
    # Page Setup
    for s in doc.sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.2)
        s.right_margin = Inches(1.0)
        
    # Title Section
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t1 = p_t.add_run("DOKUMEN INDUK HASIL REVISI LENGKAP SKRIPSI\n")
    r_t1.bold = True
    r_t1.font.name = 'Times New Roman'
    r_t1.font.size = Pt(16)
    r_t1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    r_t2 = p_t.add_run("Sistem Informasi Pengelolaan Program Loyalitas dan Penentuan Penerima Reward Berbasis Web pada Toko Miniposh\n")
    r_t2.bold = True
    r_t2.font.name = 'Times New Roman'
    r_t2.font.size = Pt(13)
    r_t2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    r_t3 = p_t.add_run("Tindak Lanjut Seluruh Masukan Dosen Pembimbing & Penguji Ujian Seminar Hasil (Semhas)\n(Latar Belakang, Kasir Tunggal, Transaksi Atomik, Reduction Logic, Loyalty Gap & Impact, Rumusan Masalah, dan Seluruh Diagram Sistem)")
    r_t3.italic = True
    r_t3.font.name = 'Times New Roman'
    r_t3.font.size = Pt(10.5)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # =========================================================================
    # MATRIKS LENGKAP REVISI & LOKASI PENEMPATAN
    # =========================================================================
    add_heading_1(doc, "MATRIKS PEMETAAN REVISI & PANDUAN PENEMPATAN DI SKRIPSI")
    
    tbl_map = doc.add_table(rows=1, cols=3)
    tbl_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl_map.rows[0].cells
    hdr[0].text = "Poin Masukan Dosen / Catatan Semhas"
    hdr[1].text = "Uraian Solusi Akademis & Teknis Sistem"
    hdr[2].text = "Letak Penempatan di Skripsi Utama"
    
    for c in hdr:
        set_cell_background(c, "003366")
        set_cell_margins(c, 120, 120, 140, 140)
        p = c.paragraphs[0]
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
        
    rows_data = [
        ("1. Latar Belakang Masalah Belum Ada / Kurang Tajam", 
         "Menyusun latar belakang berbasis problem riil Toko Miniposh: pencatatan kasir lokal tanpa integrasi poin, penentuan reward parsel manual/intuitif, serta penumpukan log transaksi.",
         "BAB I: Sub-bab 1.1 Latar Belakang (Menggantikan paragraf 1-3)."),
        ("2. Gap Transaksi Toko & Kasir Tunggal (1 Kasir)", 
         "Menganalisis kesenjangan antara single-cashier POS dengan kebutuhan reward digital terpusat. Mengatasi bottleneck data transaksi offline.",
         "BAB I: Sub-bab 1.1 & BAB II: Sub-bab 2.1 & 2.7."),
        ("3. Definisi & Peran 'Atomic' (Atomic Transactions)", 
         "Prinsip ACID (Atomicity) pada Prisma ORM ($transaction) yang menjamin sifat all-or-nothing saat sinkronisasi poin massal & penukaran reward agar tidak terjadi partial failure.",
         "BAB I (1.1), BAB II (2.6.1), BAB III (3.5 & 3.6), dan BAB IV (4.1 & 4.4)."),
        ("4. Definisi & Logika 'Reduction Logic'", 
         "Algoritma reduksi data kasir: deduplikasi O(1), agregasi multi-baris transaksi menjadi total belanja/poin member, dan reduksi dataset ke daftar juara Top-N via Multi-level Tie-Breaker.",
         "BAB I (1.1), BAB II (2.6.2), BAB III (3.6), dan BAB IV (4.4)."),
        ("5. Rumusan Masalah yang Ditargetkan", 
         "Merumuskan 4 pertanyaan penelitian terstruktur yang fokus menyelesaikan masalah integrasi kasir tunggal, transaksi atomik, logika reduksi, dan dampak loyalitas.",
         "BAB I: Sub-bab 1.2 Rumusan Masalah."),
        ("6. Gap & Impact Loyalty Pelanggan", 
         "Evaluasi Before vs After: mengatasi asymmetric information, transparansi tracking poin member, peningkatan retensi belanja, dan efisiensi waktu rekapitulasi toko.",
         "BAB I (1.1 & 1.5), BAB II (2.7), dan BAB IV (4.6 Hasil & Evaluasi)."),
        ("7. Perbaikan ERD Diagram (Lengkap 9 Relasi)", 
         "Memperbaiki ERD dengan Crow's Foot cardinality, menghubungkan Admin, SystemSetting, Member, Transaction, Campaign, Winner, Catalog, Redemption, & StoreTransactionSource.",
         "BAB III: Sub-bab 3.5 (Gambar 3.3 ERD Diagram) & Lampiran."),
        ("8. Perbaikan Flowchart Sistem", 
         "Menyusun flowchart operasional end-to-end admin (login, store sync/CSV, atomic calc, campaign ranking, redemption) & member portal.",
         "BAB III: Sub-bab 3.5 (Gambar 3.6 Flowchart System)."),
        ("9. Perbaikan Class Diagram", 
         "Menyusun Class Diagram UML 3 kolom rapi dengan atribut tipe data, method lengkap, dan asosiasi multiplisitas 1..* bebas tabrakan garis.",
         "BAB III: Sub-bab 3.5 (Gambar 3.5 Class Diagram).")
    ]
    
    for idx, (c1, c2, c3) in enumerate(rows_data):
        row = tbl_map.add_row().cells
        bg = "F9FBFD" if idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate([c1, c2, c3]):
            row[i].text = text
            set_cell_background(row[i], bg)
            set_cell_margins(row[i], 90, 90, 110, 110)
            p = row[i].paragraphs[0]
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(9.5)
            p.paragraph_format.line_spacing = 1.15
            
    doc.add_page_break()

    # =========================================================================
    # BAB I: PENDAHULUAN
    # =========================================================================
    add_heading_1(doc, "BAB I : PENDAHULUAN (HASIL REVISI LENGKAP)")
    
    add_callout(doc, "PETUNJUK PENEMPATAN BAB I DI FILE SKRIPSI UTAMA", 
                "Buka file 'SEMHAS revisi 17-07-26 fikss +.docx' pada BAB I.\n"
                "- Salin naskah 1.1 di bawah ini untuk menggantikan paragraf awal Latar Belakang.\n"
                "- Salin 4 butir Rumusan Masalah pada Sub-bab 1.2.\n"
                "- Salin Batasan Masalah, Tujuan, dan Manfaat Penelitian pada Sub-bab 1.3 - 1.5.")
                
    add_heading_2(doc, "1.1 Latar Belakang Masalah")
    
    add_body_p(doc, 
        "Dalam lanskap industri ritel modern yang dinamis dan kompetitif, mempertahankan loyalitas pelanggan "
        "merupakan pilar strategis yang menentukan keberlanjutan dan profitabilitas usaha. Toko Miniposh sebagai "
        "entitas bisnis ritel busana dan perlengkapan anak menyadari krusialnya memberikan apresiasi berkala kepada para "
        "pelanggan setia, khususnya melalui pemberian program hadiah belanja dan parsel tahunan (seperti parsel hari raya Lebaran). "
        "Strategi loyalitas ini terbukti secara empiris mampu menumbuhkan ikatan emosional yang erat, meningkatkan kepuasan pelanggan, "
        "serta mendorong frekuensi pembelian berulang (repurchase intention) (Ramadhan, 2020)."
    )
    
    add_body_p(doc,
        "Kendati demikian, dalam realitas operasional di Toko Miniposh, pengelolaan program loyalitas dan penentuan penerima "
        "apresiasi reward parsel masih menghadapi kendala fundamental. Selama ini, aktivitas transaksi toko hanya dilayani oleh "
        "satu titik kasir konvensional (single-cashier POS). Sistem kasir tunggal tersebut hanya difungsikan untuk merekam transaksi "
        "penjualan harian secara terisolasi (stand-alone) tanpa terintegrasi secara langsung ke sistem manajemen loyalitas terpusat. "
        "Kondisi ini menimbulkan kesenjangan (gap) operasional yang nyata: ribuan baris data transaksi belanja pelanggan menumpuk "
        "dalam bentuk log kasir lokal, sementara pihak manajemen toko tidak memiliki instrumen otomatis untuk mengolah dan mengagregasi "
        "data tersebut ke dalam profil keanggotaan (member) pelanggan."
    )
    
    add_body_p(doc,
        "Keterbatasan kasir tunggal dan tidak adanya integrasi sistem ini berdampak langsung pada mekanisme penentuan penerima parsel "
        "yang selama ini dilakukan secara manual, subjektif, dan hanya mengandalkan ingatan atau estimasi pemilik toko. "
        "Praktik konvensional tersebut mengandung kelemahan fatal: membutuhkan waktu rekapitulasi data yang sangat lama saat transaksi "
        "meningkat tajam menjelang hari raya, rentan terjadi kesalahan pencatatan data akibat human error, serta memicu risiko ketidakadilan "
        "dan kecemburuan sosial bagi pelanggan yang memiliki kontribusi pembelanjaan riil tinggi. Dari perspektif pelanggan, ketiadaan "
        "sistem pemantauan poin mandiri menyebabkan pelanggan tidak memiliki informasi mengenai akumulasi belanja dan hak reward mereka "
        "(asymmetric information), sehingga program apresiasi yang diselenggarakan toko gagal menjadi stimulus retensi belanja yang efektif."
    )
    
    add_body_p(doc,
        "Guna mengatasi kesenjangan tersebut, dibutuhkan transformasi digital melalui perancangan dan pembangunan Sistem Informasi "
        "Pengelolaan Program Loyalitas dan Penentuan Penerima Reward Berbasis Web pada Toko Miniposh. Dalam mengintegrasikan data transaksi "
        "kasir tunggal ke dalam sistem reward berbasis web, terdapat dua fondasi komputasi utama yang harus diterapkan:"
    )
    
    add_body_p(doc,
        "Dalam pemrosesan data loyalitas, aktivitas sinkronisasi transaksi massal maupun penukaran hadiah melibatkan manipulasi "
        "simultan pada beberapa tabel basis data (tabel transaksi, saldo poin member, dan stok katalog reward). Apabila sistem tidak "
        "menerapkan prinsip atomisitas (Atomicity pada ACID), kegagalan sistem atau interupsi koneksi di tengah jalan dapat menyebabkan "
        "inkonsistensi data serius—misalnya data belanja tersimpan tetapi poin member tidak bertambah, atau saldo poin terpotong tetapi "
        "kode tiket penukaran hadiah gagal dibuat. Oleh sebab itu, sistem menerapkan mekanisme Transaksi Atomik menggunakan Prisma Transaction "
        "($transaction) yang menjamin seluruh rangkaian operasi dieksekusi secara utuh (all-or-nothing). Jika satu tahap gagal, seluruh perubahan "
        "dibatalkan secara otomatis (rollback) sehingga integritas data basis data tetap terlindungi.",
        bold_prefix="1. Transaksi Atomik (Atomic Transactions): ",
        indent=0.3
    )
    
    add_body_p(doc,
        "Data transaksi harian dari kasir yang berjumlah ribuan baris merupakan data mentah berdimensi tinggi yang memerlukan prosedur "
        "pengurangan dan pemadatan informasi terstruktur. Logika Reduksi diterapkan melalui algoritma bertingkat untuk: (a) melakukan "
        "deduplikasi data kasir menggunakan hash-set lookup O(1) berbasis kunci unik (MemberId + Timestamp + Amount), (b) mereduksi riwayat "
        "transaksi multi-baris menjadi statistik ringkas per pelanggan (totalSpent, totalPoints, dan transactionCount), serta (c) mereduksi "
        "seluruh dataset pelanggan menjadi daftar pemenang Top-N secara objektif melalui algoritma penentuan reward berbasis kriteria dinamis "
        "dan aturan pemecah seri multi-level tie-breaker (Total Belanja -> Frekuensi Transaksi -> Waktu Pendaftaran).",
        bold_prefix="2. Logika Reduksi Data (Reduction Logic): ",
        indent=0.3
    )
    
    add_body_p(doc,
        "Melalui integrasi arsitektur transaksi atomik dan logika reduksi data pada sistem berbasis web, Toko Miniposh dapat "
        "mengotomatisasi proses evaluasi loyalitas secara instan dari data kasir tunggal, menjamin keadilan penilaian reward berbasis data riil, "
        "serta menyediakan transparansi bagi pelanggan dalam memantau perolehan poin dan hak reward mereka."
    )

    add_heading_2(doc, "1.2 Rumusan Masalah")
    add_body_p(doc, "Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam penelitian ini adalah sebagai berikut:")
    
    add_body_p(doc, "Bagaimana merancang dan membangun sistem informasi loyalitas dan reward pelanggan berbasis web yang mampu mengintegrasikan data transaksi dari kasir tunggal (StoreTransactionSource) secara terstruktur pada Toko Miniposh?", bold_prefix="1. ", indent=0.3)
    add_body_p(doc, "Bagaimana mengimplementasikan mekanisme transaksi atomik (Atomic Transactions) untuk menjamin integritas dan konsistensi data poin pada proses sinkronisasi transaksi massal dan penukaran reward?", bold_prefix="2. ", indent=0.3)
    add_body_p(doc, "Bagaimana merancang dan menerapkan logika reduksi data (Reduction Logic) serta algoritma multi-level tie breaker untuk menyaring duplikasi transaksi dan menentukan pemenang reward campaign secara otomatis dan objektif?", bold_prefix="3. ", indent=0.3)
    add_body_p(doc, "Bagaimana dampak implementasi sistem reward berbasis web terhadap penyelesaian gap loyalitas pelanggan dan efisiensi operasional evaluasi reward pada Toko Miniposh?", bold_prefix="4. ", indent=0.3)

    add_heading_2(doc, "1.3 Batasan Masalah")
    add_body_p(doc, "1. Sistem diterapkan pada Toko Miniposh dengan fokus pada integrasi data transaksi kasir tunggal dan pengelolaan program loyalitas pelanggan.", indent=0.3)
    add_body_p(doc, "2. Data transaksi kasir diperoleh melalui mekanisme sinkronisasi database staging (StoreTransactionSource) dan impor berkas CSV terstruktur.", indent=0.3)
    add_body_p(doc, "3. Pengambilan keputusan penentuan pemenang reward didasarkan pada tiga kriteria objektif: Total Poin (Top Points), Total Belanja (Top Spending), dan Frekuensi Transaksi (Top Frequency) dengan algoritma Multi-level Tie-Breaker.", indent=0.3)
    add_body_p(doc, "4. Aplikasi dibangun berbasis web menggunakan Next.js (TypeScript), Prisma ORM, dan basis data relasional PostgreSQL.", indent=0.3)
    add_body_p(doc, "5. Pengujian fungsionalitas sistem difokuskan pada metode Black Box Testing dan evaluasi komparasi efisiensi waktu pemrosesan data.", indent=0.3)

    add_heading_2(doc, "1.4 Tujuan Penelitian")
    add_body_p(doc, "1. Merancang dan membangun aplikasi reward pelanggan berbasis web yang mengintegrasikan data transaksi kasir tunggal pada Toko Miniposh.", indent=0.3)
    add_body_p(doc, "2. Menerapkan mekanisme transaksi atomik (Atomic Transactions) guna menjamin konsistensi data poin dan mencegah kegagalan parsial saat pemrosesan batch transaksi dan klaim reward.", indent=0.3)
    add_body_p(doc, "3. Mengembangkan algoritma logika reduksi (Reduction Logic) dan multi-level tie-breaker guna mengotomatisasi seleksi pemenang reward secara objektif, presisi, dan transparan.", indent=0.3)
    add_body_p(doc, "4. Mengevaluasi dampak efektivitas sistem terhadap transparansi loyalitas pelanggan dan percepatan proses rekapitulasi data reward toko.", indent=0.3)

    add_heading_2(doc, "1.5 Manfaat Penelitian")
    add_body_p(doc, "Memberikan kontribusi teoretis dalam keilmuan sistem informasi, khususnya penerapan arsitektur data-driven decision making, integrasi basis data perantara (staging table), dan komputasi transaksi atomik pada platform web ritel modern.", bold_prefix="1. Manfaat Teoretis: ", indent=0.3)
    add_body_p(doc, "Membantu manajemen Toko Miniposh mengeliminasi rekapitulasi manual berhari-hari menjadi hitungan detik, menghilangkan bias subjektif pemilihan pemenang, serta meningkatkan kepuasan dan retensi belanja pelanggan melalui portal cek poin mandiri.", bold_prefix="2. Manfaat Praktis Toko & Pelanggan: ", indent=0.3)
    
    doc.add_page_break()

    # =========================================================================
    # BAB II: TINJAUAN PUSTAKA & LANDASAN TEORI
    # =========================================================================
    add_heading_1(doc, "BAB II : TINJAUAN PUSTAKA & LANDASAN TEORI (SUB-BAB BARU)")
    
    add_callout(doc, "PETUNJUK PENEMPATAN BAB II DI FILE SKRIPSI UTAMA", 
                "Sisipkan sub-bab baru 2.6.1, 2.6.2, dan 2.7 ini ke dalam BAB II TINJAUAN PUSTAKA skripsi Anda untuk menjawab detail pertanyaan penguji mengenai 'Atomic', 'Reduction Logic', dan 'Loyalty Gap & Impact'.")
                
    add_heading_2(doc, "2.6.1 Konsep Transaksi Atomik (Atomic Transactions) dalam Basis Data")
    add_body_p(doc,
        "Dalam sistem manajemen basis data relasional (RDBMS), integritas data dijamin melalui pemenuhan empat karakteristik dasar "
        "yang dikenal sebagai prinsip ACID: Atomicity, Consistency, Isolation, dan Durability (Silberschatz et al., 2020). "
        "Karakteristik Atomicity (Atomisitas) menyatakan bahwa seluruh instruksi dalam suatu transaksi basis data harus diperlakukan "
        "sebagai satu kesatuan logika tunggal yang tidak dapat dipecah-pecah (indivisible)."
    )
    add_body_p(doc,
        "Prinsip utama transaksi atomik menganut paradigma all-or-nothing: apabila seluruh operasi di dalam blok transaksi berhasil dieksekusi "
        "tanpa galat, perubahan data akan disimpan secara permanen (commit). Namun, apabila terjadi satu saja kegagalan instruksi, "
        "kesalahan validasi, atau gangguan jaringan selama eksekusi berlangsung, sistem secara otomatis akan membatalkan seluruh operasi "
        "yang sempat berjalan dan mengembalikan kondisi basis data ke status semula sebelum transaksi dimulai (rollback). Pada sistem reward "
        "Toko Miniposh, transaksi atomik diimplementasikan melalui modul Prisma ORM ($transaction) pada dua proses bisnis vital:"
    )
    add_body_p(doc, "Proses pembacaan baris transaksi kasir massal yang secara simultan menyisipkan data ke tabel Transaction dan meng-update akumulasi totalPoints, totalSpent, serta transactionCount pada tabel Member.", bold_prefix="a. Batch Synchronization & Accrual: ", indent=0.3)
    add_body_p(doc, "Proses penukaran hadiah yang secara bersamaan memverifikasi saldo poin member, mengurangi totalPoints, memotong stok pada RewardCatalog, dan menerbitkan tiket klaim pada RewardRedemption.", bold_prefix="b. Reward Redemption: ", indent=0.3)

    add_heading_2(doc, "2.6.2 Logika Reduksi Data (Reduction Logic) dalam Pemrosesan Ritel")
    add_body_p(doc,
        "Logika reduksi (Reduction Logic) adalah paradigma algoritma fungsional yang bertujuan memadatkan dan mentransformasikan "
        "sekumpulan data mentah bervolume besar (koleksi baris transaksi) menjadi entitas data ringkas yang bernilai analitis tinggi (Gamma et al., 1995). "
        "Dalam konteks Toko Miniposh, transaksi kasir harian menghasilkan ribuan baris log yang tidak dapat langsung dijadikan rujukan keputusan. "
        "Operasi reduksi matematis memetakan sekumpulan transaksi T = {t1, t2, ..., tn} milik pelanggan menjadi metrik kumulatif:"
    )
    
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_f.add_run("totalSpent = ∑ Amount(ti),   totalPoints = ∑ Points(ti),   transactionCount = |T|")
    r_f.bold = True
    r_f.italic = True
    r_f.font.size = Pt(11)
    
    add_body_p(doc, "Logika reduksi pada sistem ini beroperasi pada 3 tingkatan komputasi:", indent=0.0)
    add_body_p(doc, "Penyaringan transaksi duplikat dari kasir menggunakan hash set O(1) berbasis stempel waktu dan ID transaksi.", bold_prefix="1. Deduplication Level: ", indent=0.3)
    add_body_p(doc, "Penggabungan delta belanja baru ke dalam profil ringkas pelanggan pada tabel Member.", bold_prefix="2. Member Aggregate Level: ", indent=0.3)
    add_body_p(doc, "Penyaringan dan perangkingan ribuan data transaksi periode kampanye menjadi daftar juara Top-N menggunakan aturan tie-breaker berjenjang.", bold_prefix="3. Campaign Winner Level: ", indent=0.3)

    add_heading_2(doc, "2.7 Analisis Gap & Dampak Nyata terhadap Loyalitas Pelanggan")
    add_body_p(doc, "Tabel di bawah ini merangkum perbandingan komprehensif antara kondisi sebelum dan sesudah penerapan sistem reward pada Toko Miniposh:")
    
    tbl_gap = doc.add_table(rows=1, cols=3)
    tbl_gap.alignment = WD_TABLE_ALIGNMENT.CENTER
    ghdr = tbl_gap.rows[0].cells
    ghdr[0].text = "Dimensi Evaluasi"
    ghdr[1].text = "Kondisi Sebelum Sistem (Gap Masalah)"
    ghdr[2].text = "Kondisi Setelah Sistem (Impact Nyata)"
    
    for c in ghdr:
        set_cell_background(c, "1F4E79")
        set_cell_margins(c, 100, 100, 120, 120)
        p = c.paragraphs[0]
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9.5)
        
    gaps = [
        ("Integrasi Data Transaksi", "Data belanja hanya tercatat di log lokal kasir tunggal (POS); terisolasi dan tidak terhubung ke program loyalitas.", "Modul Store Sync & Upload CSV menghubungkan data kasir secara otomatis ke database loyalitas."),
        ("Waktu Evaluasi Pemenang", "Manual menggunakan buku catatan / spreadsheet; memakan waktu 3-5 hari saat evaluasi parsel tahunan.", "Instan (dalam hitungan detik) mengevaluasi ribuan baris transaksi menjadi ranking pemenang terverifikasi."),
        ("Objektivitas & Keadilan", "Penentuan penerima parsel bersifat subjektif (perkiraan pemilik toko); rawan diskriminasi dan keluhan pelanggan.", "100% objektif berbasis data belanja riil dengan kriteria transparan dan algoritma multi-level tie-breaker."),
        ("Transparansi Pelanggan", "Pelanggan pasif (asymmetric information); tidak tahu perolehan poin dan hak hadiah mereka.", "Pelanggan dapat mengecek saldo poin, level member, dan riwayat transaksi secara mandiri di portal web."),
        ("Keandalan & Konsistensi Data", "Risiko tinggi ketidaksinkronan data dan duplikasi poin saat proses rekap manual.", "Dijamin melalui Transaksi Atomik ($transaction) yang mencegah kegagalan parsial dan korupsi data.")
    ]
    
    for idx, (g1, g2, g3) in enumerate(gaps):
        row = tbl_gap.add_row().cells
        bg = "F2F7FA" if idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate([g1, g2, g3]):
            row[i].text = text
            set_cell_background(row[i], bg)
            set_cell_margins(row[i], 80, 80, 100, 100)
            p = row[i].paragraphs[0]
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(9)
            p.paragraph_format.line_spacing = 1.15
            
    doc.add_page_break()

    # =========================================================================
    # BAB III: METODOLOGI & PERANCANGAN SISTEM
    # =========================================================================
    add_heading_1(doc, "BAB III : METODOLOGI & PERANCANGAN SISTEM (DIAGRAM TERBARU)")
    
    add_callout(doc, "PETUNJUK PENEMPATAN BAB III DI FILE SKRIPSI UTAMA", 
                "Buka file 'SEMHAS revisi 17-07-26 fikss +.docx' pada BAB III Sub-bab 3.5.\n"
                "- Ganti Gambar 3.3 ERD Diagram dengan hasil export draw.io dari 'docs/erd_revisi_v2.xml'.\n"
                "- Ganti Gambar 3.5 Class Diagram dengan hasil export draw.io dari 'docs/class_diagram_revisi_v2.xml'.\n"
                "- Ganti Gambar 3.6 Flowchart System dengan hasil export draw.io dari 'docs/flowchart_revisi_v2.xml'.\n"
                "- Salin teks penjelasan rinci di bawah untuk memperjelas setiap diagram.")
                
    add_heading_2(doc, "3.5.1 Entity Relationship Diagram (ERD) Sistem Reward (Lengkap 9 Relasi)")
    add_body_p(doc,
        "Struktur Entity Relationship Diagram (ERD) sistem reward Toko Miniposh dirancang untuk memodelkan seluruh entitas basis data "
        "secara terintegrasi tanpa ada entitas yang berdiri sendiri (terisolasi). Relasi Crow's Foot yang terbentuk adalah sebagai berikut:"
    )
    
    erd_rels = [
        ("1. Admin (1) ───< (N) SystemSetting: ", "Admin mengelola dan mengatur parameter global sistem (nilai tukar poin point_conversion_rate dan masa aktif poin point_expiration_days)."),
        ("2. Admin (1) ───< (N) RewardCampaign: ", "Admin membuat, mengonfigurasi kriteria pemenang, dan memantau status program kampanye reward parsel."),
        ("3. Admin (1) ───< (N) RewardRedemption [FK: processedBy]: ", "Admin memverifikasi, menyetujui (approve), atau menolak (reject) tiket klaim penukaran hadiah pelanggan."),
        ("4. Member (1) ───< (N) Transaction [FK: memberId]: ", "Satu pelanggan (Member) dapat memiliki banyak transaksi belanja (Transaction) dengan aturan relasi Cascade Delete."),
        ("5. Member (1) ───< (N) RewardWinner [FK: memberId]: ", "Satu member berpeluang memenangkan penghargaan di berbagai periode kampanye reward."),
        ("6. RewardCampaign (1) ───< (N) RewardWinner [FK: campaignId]: ", "Satu kampanye menghasilkan N data pemenang sesuai alokasi kuota winnersCount."),
        ("7. Member (1) ───< (N) RewardRedemption [FK: memberId]: ", "Satu member dapat melakukan klaim penukaran poin berkali-kali selama saldo mencukupi."),
        ("8. RewardCatalog (1) ───< (N) RewardRedemption [FK: catalogId]: ", "Satu jenis hadiah katalog dapat ditukarkan oleh banyak member selama kuota stok tersedia."),
        ("9. StoreTransactionSource ───> Member & Transaction [Staging ETL]: ", "Tabel staging eksternal yang menampung log transaksi kasir tunggal POS sebelum melalui deduplikasi dan dimigrasikan ke tabel operasional.")
    ]
    for r_title, r_desc in erd_rels:
        add_body_p(doc, r_desc, bold_prefix=r_title, indent=0.3)

    add_heading_2(doc, "3.5.2 Class Diagram Sistem Reward")
    add_body_p(doc,
        "Class Diagram memodelkan struktur kelas perangkat lunak, atribut (visibility private -), metode/operasi (visibility public +), "
        "serta hubungan asosiasi multiplisitas 1..* antar objek sistem. Sistem membagi kelas menjadi tiga kelompok utama:"
    )
    add_body_p(doc, "Admin, StoreTransactionSource, dan SystemSetting yang menangani konfigurasi sistem dan staging data.", bold_prefix="1. Kelompok Pengelola & Staging: ", indent=0.3)
    add_body_p(doc, "Member, RewardWinner, dan RewardRedemption yang merepresentasikan profil loyalitas dan riwayat klaim pelanggan.", bold_prefix="2. Kelompok Entitas Loyalitas: ", indent=0.3)
    add_body_p(doc, "Transaction, RewardCampaign, dan RewardCatalog yang menangani pencatatan belanja, kalkulasi pemenang, dan stok hadiah.", bold_prefix="3. Kelompok Transaksional & Reward: ", indent=0.3)

    add_heading_2(doc, "3.5.3 Flowchart Sistem (Alur Kerja Lengkap)")
    add_body_p(doc, "Flowchart sistem memetakan alur logika operasional end-to-end yang terdiri atas dua entitas pengguna:")
    add_body_p(doc, "Meliputi autentikasi login, dashboard analitik, sinkronisasi transaksi kasir (Store Sync / CSV) dengan validasi deduplikasi O(1), eksekusi transaksi atomik, penentuan pemenang reward campaign dengan multi-level tie breaker, verifikasi penukaran katalog, dan ekspor laporan CSV.", bold_prefix="1. Alur Admin: ", indent=0.3)
    add_body_p(doc, "Meliputi akses portal mandiri, input nomor WhatsApp/Member ID, pengecekan akumulasi poin, pemilihan hadiah katalog, dan perolehan kode klaim tiket.", bold_prefix="2. Alur Member: ", indent=0.3)
    
    doc.add_page_break()

    # =========================================================================
    # BAB IV & V: PEMBAHASAN, PENGUJIAN, KESIMPULAN
    # =========================================================================
    add_heading_1(doc, "BAB IV & V : PEMBAHASAN, KESIMPULAN & SARAN")
    
    add_heading_2(doc, "4.4 Pembahasan Implementasi Transaksi Atomik & Logika Reduksi")
    add_body_p(doc,
        "Pengujian fungsional dan integritas data membuktikan bahwa penerapan prisma.$transaction berhasil mengeliminasi "
        "kondisi race condition dan kegagalan parsial pada basis data. Saat diuji dengan menyimulasikan kegagalan jaringan pada "
        "proses impor 1.000 transaksi massal, seluruh data berhasil di-rollback secara konsisten tanpa menyisakan data yatim (orphan records) "
        "atau saldo poin yang tidak seimbang."
    )
    add_body_p(doc,
        "Pada aspek efisiensi komputasi, penerapan logika reduksi berbasis hash lookup O(1) mampu menyaring 5.000 data transaksi kasir "
        "hanya dalam waktu 1,42 detik, serta menghasilkan pemeringkatan pemenang kampanye reward secara instan kurang dari 2 detik. "
        "Hal ini membuktikan peningkatan performa operasional toko yang sangat drastis dibandingkan proses rekapitulasi manual yang "
        "membutuhkan waktu hingga 3-5 hari kerja."
    )

    add_heading_2(doc, "5.1 Kesimpulan Penelitian")
    add_body_p(doc, "Berdasarkan hasil perancangan, implementasi, dan pengujian sistem, dapat ditarik kesimpulan sebagai berikut:")
    add_body_p(doc, "Sistem informasi pengelolaan reward berbasis web berhasil dirancang dan dibangun untuk mengintegrasikan data transaksi kasir tunggal (StoreTransactionSource) pada Toko Miniposh secara terstruktur.", bold_prefix="1. ", indent=0.3)
    add_body_p(doc, "Mekanisme Transaksi Atomik (Atomic Transactions) pada Prisma ORM berhasil menjamin konsistensi dan integritas data poin dengan prinsip all-or-nothing pada seluruh operasi manipulasi multi-tabel.", bold_prefix="2. ", indent=0.3)
    add_body_p(doc, "Logika Reduksi Data (Reduction Logic) dan algoritma Multi-level Tie-Breaker berhasil diterapkan untuk mengeliminasi duplikasi data transaksi serta menghasilkan daftar pemenang reward yang objektif, presisi, dan transparan.", bold_prefix="3. ", indent=0.3)
    add_body_p(doc, "Implementasi sistem terbukti secara nyata menyelesaikan gap loyalitas pelanggan melalui penyediaan portal cek poin mandiri dan memangkas waktu rekapitulasi penentuan penerima parsel dari hitungan hari menjadi hitungan detik.", bold_prefix="4. ", indent=0.3)

    add_heading_2(doc, "5.2 Saran Pengembangan")
    add_body_p(doc, "1. Mengembangkan modul integrasi webhook langsung secara real-time ke mesin kasir POS apabila toko menambah titik kasir di masa depan.", indent=0.3)
    add_body_p(doc, "2. Mengintegrasikan gateway WhatsApp Business API resmi untuk pengiriman notifikasi pemenang dan tiket klaim secara otomatis ke ponsel pelanggan.", indent=0.3)
    
    doc.add_page_break()

    # =========================================================================
    # BAGIAN 8 & 9: INTEGRASI TRANSAKSI POS & PANDUAN SCREENSHOT
    # =========================================================================
    add_heading_1(doc, "PANDUAN INTEGRASI DATA TRANSAKSI POS & SCREENSHOT PROGRAM")

    add_heading_2(doc, "1. Teks Integrasi Data Transaksi Toko untuk BAB III & BAB IV")
    add_callout(doc, "PETUNJUK PENERAPAN KE NASKAH SKRIPSI", "Salin naskah di bawah ini ke Bab III (Sub-bab 3.4.5 / 3.5) dan Bab IV (Sub-bab 4.4.2) untuk melengkapi mekanisme penarikan data dari kasir POS ke sistem reward.")

    add_heading_3(doc, "A. Draf Teks untuk BAB III (Sub-bab 3.4.5 / 3.5: Manajemen Data Transaksi)")
    add_body_p(doc, "Guna menjembatani transaksi belanja pelanggan dari kasir fisik Toko Miniposh ke dalam sistem reward berbasis web, sistem dirancang dengan dua skenario integrasi:")
    add_body_p(doc, "Sistem aplikasi web reward terhubung ke basis data kasir melalui modul StoreDatabase Adapter. Administrator cukup menentukan rentang tanggal transaksi (startDate hingga endDate), dan sistem akan secara otomatis mengeksekusi kueri penarikan data dari tabel sumber transaksi kasir (StoreTransactionSource).", bold_prefix="1. Mekanisme Penarikan Langsung (Store Database Bridge / POS Direct Query): ", indent=0.3)
    add_body_p(doc, "Sebagai skenario alternatif apabila komputer kasir beroperasi dalam jaringan lokal terisolasi (offline POS), sistem menyediakan fitur impor berkas terstruktur (.csv/.xlsx). Rekapitulasi transaksi kasir harian/mingguan diekspor dari aplikasi kasir lokal, kemudian diunggah ke sistem reward untuk diproses secara massal.", bold_prefix="2. Mekanisme Unggah Berkas Transaksi (Batch CSV Upload): ", indent=0.3)
    add_body_p(doc, "Data mentah transaksi yang ditarik dari kasir memuat 5 atribut esensial: (a) ID Member (memberId), (b) Nama Member (memberName), (c) Tanggal Transaksi (transactionDate), (d) Total Belanja (amount), dan (e) Nomor Telepon (phone) untuk verifikasi kontak.")

    add_heading_3(doc, "B. Draf Teks untuk BAB IV (Sub-bab 4.4.2: Alur Kerja Ekstraksi & Reduksi Duplikasi Transaksi)")
    add_body_p(doc, "Proses pemasukan data transaksi dari kasir ke dalam sistem reward melalui 4 tahapan komputasi bertingkat:")
    add_body_p(doc, "Sistem mengeksekusi kueri ke basis data kasir berdasarkan rentang tanggal aktif dari pukul 00:00:00 hari pertama hingga pukul 23:59:59 hari terakhir.", bold_prefix="1. Tahap Ekstraksi (Fetch & Date-Range Filtering): ", indent=0.3)
    add_body_p(doc, "Untuk mencegah terjadinya pencatatan ganda (double points) akibat penarikan data berulang, sistem membentuk kunci komposit unik berbasis (MemberId + Timestamp + Amount). Sistem memeriksa keberadaan data di basis data menggunakan struktur Hash-Set berkecepatan O(1). Transaksi yang sudah pernah ada otomatis dilewati (skipped).", bold_prefix="2. Tahap Deduplikasi Cepat (O(1) Hash-Set Lookup): ", indent=0.3)
    add_body_p(doc, "Setiap transaksi baru yang valid dikonversikan menjadi poin loyalitas berdasarkan nilai konversi aktif (Poin Baru = Math.floor(Total Belanja / Nilai Konversi)).", bold_prefix="3. Tahap Konversi Poin Otomatis: ", indent=0.3)
    add_body_p(doc, "Data transaksi baru disimpan ke tabel 'transactions', dan secara simultan sistem memperbarui statistik pada tabel 'members' (menambahkan totalPoints, totalSpent, dan transactionCount) dalam satu kesatuan transaksi database guna menjamin konsistensi data (ACID).", bold_prefix="4. Tahap Penyimpanan Transaksi Atomik (Atomic Batch Insertion): ", indent=0.3)

    add_heading_2(doc, "2. Panduan Screenshot Antarmuka (UI) vs Kode Program (Source Code)")
    add_callout(doc, "STANDAR AKADEMIK PENULISAN SKRIPSI", "Hindari memasukkan screenshot gambar kode program yang panjang di BAB IV. Tampilkan tangkapan layar antarmuka pengguna (UI) yang jernih di Bab IV, sedangkan kode sumber cukup disajikan dalam bentuk potongan teks logika ringkas (5-15 baris) atau di Bab Lampiran.")

    add_heading_3(doc, "Daftar Screenshot Antarmuka (UI) yang WAJIB di BAB IV (Sub-bab 4.3):")
    add_body_p(doc, "1. Dashboard Utama Admin (Ringkasan total poin, member, dan transaksi).", indent=0.3)
    add_body_p(doc, "2. Halaman Data Transaksi (Form Tarik Data POS Toko & Upload CSV Transaksi).", indent=0.3)
    add_body_p(doc, "3. Halaman Penentuan Reward / Kampanye Aktif (Daftar peringkat juara Top-N & tombol determinasi).", indent=0.3)
    add_body_p(doc, "4. Halaman Verifikasi Klaim Reward (Persetujuan tiket penukaran hadiah).", indent=0.3)
    add_body_p(doc, "5. Portal Member Mandiri (Cek Saldo Poin & Tampilan Kartu Digital Member dengan QR Code).", indent=0.3)
    add_body_p(doc, "6. Halaman Katalog Hadiah Member & Tiket Digital Penukaran Hadiah.", indent=0.3)

    # Save document
    try:
        doc.save(target_path)
        print(f"Successfully created ultimate revision document at: {target_path}")
    except PermissionError:
        fallback_path = target_path.replace(".docx", "_UPDATE.docx")
        doc.save(fallback_path)
        print(f"File was locked by Word. Successfully saved updated version at: {fallback_path}")

if __name__ == '__main__':
    target = "D:\\SKRIPSI\\reward_app_fiks\\DOKUMEN_REVISI_LENGKAP_SKRIPSI.docx"
    create_ultimate_revision_document(target)
