# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    return h

def add_callout(doc, title, text, bg_color="F0F4F8", border_color="1F4E79"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"📍 {title}\n")
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(10)
    r_text.italic = True
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def create_full_revision_docx():
    doc = docx.Document()
    
    # Page Setup - A4 with standard 1 inch margins
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)
    
    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("DOKUMEN LENGKAP REVISI SKRIPSI & PANDUAN INTEGRASI\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Sistem Informasi Pengelolaan Program Loyalitas dan Penentuan Penerima Reward Berbasis Web pada Toko Miniposh\n(Hasil Tindak Lanjut Masukan Ujian Seminar Hasil / Dosen Pembimbing)")
    r_sub.italic = True
    r_sub.font.size = Pt(11)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # =========================================================================
    # DAFTAR PEMETAAN REVISI
    # =========================================================================
    add_styled_heading(doc, "RINGKASAN POIN REVISI & LOKASI PENEMPATAN", level=1)
    
    table_map = doc.add_table(rows=1, cols=3)
    table_map.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_map.rows[0].cells
    hdr[0].text = "No & Poin Revisi Dosen"
    hdr[1].text = "Uraian Perbaikan yang Dilakukan"
    hdr[2].text = "Lokasi Penempatan di Skripsi"
    
    for c in hdr:
        set_cell_background(c, "003366")
        set_cell_margins(c, 120, 120, 150, 150)
        p = c.paragraphs[0]
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
        
    data_map = [
        ("1. Latar Belakang Masalah Belum Ada / Kurang Tajam", 
         "Menyusun latar belakang berbasis masalah riil toko: kasir tunggal (single cashier), data transaksi terisolasi di POS lokal, penentuan reward manual & subjektif, serta risiko ketidakpuasan pelanggan.",
         "BAB I: Sub-bab 1.1 Latar Belakang (Menggantikan/Memperkuat paragraf 1-3)"),
        ("2. Gap Transaksi Toko & Kasir Tunggal (1 Kasir)", 
         "Menjelaskan gap integrasi antara database POS lokal (StoreTransactionSource) dengan sistem loyalitas web. Menganalisis kendala kasir tunggal sebagai single point of transaction.",
         "BAB I: Sub-bab 1.1 & BAB II: Sub-bab 2.1 (Tinjauan Arsitektur Integrasi Data)"),
        ("3. Definisi & Peran 'Atomic' (Atomicity)", 
         "Menjelaskan konsep Atomic Transactions (ACID) pada Prisma ORM ($transaction) pada sinkronisasi transaksi massal dan penukaran reward agar tidak terjadi partial data corruption.",
         "BAB I (Latar Belakang), BAB II (Landasan Teori Baru), dan BAB III (Sub-bab 3.5 & 3.6)"),
        ("4. Definisi & Logika 'Reduction Logic'", 
         "Menjelaskan algoritma reduksi data: (a) deduplikasi data kasir O(1), (b) agregasi multi-baris transaksi menjadi metrik member, (c) reduksi dataset ke ranking Top-N dengan Multi-level Tie Breaker.",
         "BAB I (Latar Belakang), BAB II (Landasan Teori), BAB III (Sub-bab 3.6 Mekanisme Perhitungan)"),
        ("5. Rumusan Masalah yang Ditargetkan", 
         "Menyusun 4 rumusan masalah tajam: integrasi kasir tunggal, implementasi transaksi atomik & logika reduksi, penentuan pemenang objektif, dan evaluasi dampak loyalitas.",
         "BAB I: Sub-bab 1.2 Rumusan Masalah"),
        ("6. Gap & Impact Loyalty Pelanggan", 
         "Menganalisis perbandingan kondisi Sebelum vs Sesudah sistem: transparansi poin, gamifikasi reward, retensi belanja, dan efisiensi rekapitulasi toko.",
         "BAB I: Sub-bab 1.1 & Sub-bab 1.5 Manfaat, serta BAB IV: Sub-bab 4.6 Hasil dan Pembahasan"),
        ("7. Perbaikan ERD Diagram", 
         "Memperbaiki ERD dengan relasi yang benar, kardinalitas Crow's Foot (1-to-many) rapi, tipe data presisi, dan penjelasan peran tabel staging StoreTransactionSource.",
         "BAB III: Sub-bab 3.5 (Gambar 3.3 ERD Diagram) & Lampiran"),
        ("8. Perbaikan Flowchart Sistem", 
         "Menyusun flowchart komprehensif: login, percabangan kelola transaksi (CSV / Store Sync), atomic calculation, penentuan reward campaign, dan klaim hadiah.",
         "BAB III: Sub-bab 3.5 (Gambar 3.6 Flowchart System) & Lampiran"),
        ("9. Perbaikan Class Diagram", 
         "Menyusun class diagram terstruktur dengan seluruh atribut model, operasi/method service, serta relasi asosiasi dan multiplisitas yang akurat.",
         "BAB III: Sub-bab 3.5 (Gambar 3.5 Class Diagram) & Lampiran")
    ]
    
    for row_idx, (r1, r2, r3) in enumerate(data_map):
        row = table_map.add_row().cells
        bg = "F9FBFD" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate([r1, r2, r3]):
            row[i].text = text
            set_cell_background(row[i], bg)
            set_cell_margins(row[i], 100, 100, 120, 120)
            p = row[i].paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            p.paragraph_format.line_spacing = 1.15
            
    doc.add_page_break()

    # =========================================================================
    # BAGIAN 1: REVISI BAB I (PENDAHULUAN)
    # =========================================================================
    add_styled_heading(doc, "BAGIAN 1: REVISI BAB I (PENDAHULUAN)", level=1)
    
    add_callout(doc, "PETUNJUK PENEMPATAN DI FILE WORD SKRIPSI (BAB I)", 
                "Buka file skripsi Anda pada 'BAB I PENDAHULUAN'.\n"
                "1. Gantikan teks pada '1.1 Latar Belakang' dengan naskah Latar Belakang Lengkap di bawah ini (yang telah memuat gap kasir tunggal, urgensi loyalitas, atomic transaction, dan reduction logic).\n"
                "2. Gantikan butir-butir pada '1.2 Rumusan Masalah' dengan 4 rumusan masalah baru di bawah ini.\n"
                "3. Tambahkan butir Manfaat Praktis dan Teoretis pada '1.5 Manfaat Penelitian'.")
                
    add_styled_heading(doc, "1.1 Latar Belakang (Naskah Hasil Revisi Lengkap)", level=2)
    
    p = doc.add_paragraph(
        "Dalam lanskap bisnis ritel modern yang semakin kompetitif, mempertahankan loyalitas pelanggan merupakan faktor strategis "
        "yang menentukan keberlangsungan dan profitabilitas usaha. Toko Miniposh sebagai salah satu entitas bisnis ritel busana dan perlengkapan "
        "anak menyadari pentingnya memberikan apresiasi berkala kepada para pelanggan setianya, khususnya melalui pemberian reward belanja dan "
        "parsel tahunan (seperti parsel hari raya Lebaran). Strategi loyalitas ini terbukti efektif dalam menumbuhkan hubungan emosional yang erat "
        "serta mendorong pembelian berulang (repurchase intention) (Ramadhan, 2020)."
    )
    
    p = doc.add_paragraph(
        "Namun, dalam praktik operasionalnya di Toko Miniposh, pengelolaan program loyalitas dan penentuan penerima apresiasi reward "
        "masih menghadapi kendala mendasar. Selama ini, pencatatan transaksi toko hanya mengandalkan satu titik kasir konvensional (single-cashier POS). "
        "Sistem kasir tunggal ini hanya berfungsi mencatat penjualan harian secara terisolasi tanpa terhubung langsung ke sistem database reward terpadu. "
        "Akibatnya, timbul kesenjangan (gap) operasional yang signifikan: ribuan riwayat transaksi pelanggan menumpuk dalam format log penjualan lokal, "
        "sementara pihak pengelola toko tidak memiliki sarana otomatis untuk mengagregasi data belanja tersebut ke dalam profil loyalitas masing-masing pelanggan."
    )
    
    p = doc.add_paragraph(
        "Kondisi kasir tunggal dan ketiadaan integrasi sistem ini berdampak langsung pada proses penentuan penerima reward parsel yang dilakukan secara manual, "
        "subjektif, dan hanya mengandalkan intuisi atau perkiraan pemilik toko. Mekanisme subjektif tersebut memiliki kelemahan kritis: rawan memicu kekeliruan "
        "pencatatan, membutuhkan waktu rekapitulasi yang sangat lama saat volume transaksi tinggi, serta menimbulkan potensi ketidakadilan bagi pelanggan yang "
        "secara riil memiliki kontribusi transaksi besar. Dari sudut pandang pelanggan, ketiadaan sistem reward yang transparan mengakibatkan pelanggan tidak "
        "mengetahui perolehan poin dan hak reward mereka (asymmetric information), sehingga efektivitas program loyalitas dalam meningkatkan retensi belanja "
        "menjadi tidak optimal."
    )
    
    p = doc.add_paragraph(
        "Untuk mengatasi permasalahan tersebut, dibutuhkan transformasi digital berupa sistem informasi reward pelanggan berbasis web yang mengintegrasikan "
        "data transaksi kasir dengan manajemen loyalitas secara otomatis dan terstruktur. Dalam merancang arsitektur sistem ini, terdapat dua tantangan teknis utama "
        "yang harus diselesaikan:"
    )
    
    p_num1 = doc.add_paragraph()
    p_num1.paragraph_format.left_indent = Inches(0.5)
    r1 = p_num1.add_run("1. Transaksi Atomik (Atomic Transactions): ")
    r1.bold = True
    p_num1.add_run(
        "Dalam pengolahan data loyalitas, setiap proses sinkronisasi transaksi massal maupun penukaran reward (redemption) melibatkan manipulasi simultan "
        "pada multi-tabel (tabel transaksi, saldo poin member, dan stok katalog reward). Jika sistem tidak menerapkan prinsip atomisitas (Atomicity - ACID), "
        "kegagalan parsial di tengah eksekusi dapat menyebabkan inkonsistensi data parah (seperti data transaksi tersimpan namun poin member gagal bertambah, "
        "atau poin terpotong namun tiket penukaran reward gagal dibuat). Oleh karena itu, sistem harus menerapkan mekanisme transaksi atomik pada lapisan basis data "
        "menggunakan Prisma Transaction ($transaction) yang menjamin seluruh rangkaian operasi berhasil secara utuh atau dibatalkan sepenuhnya (rollback)."
    )
    
    p_num2 = doc.add_paragraph()
    p_num2.paragraph_format.left_indent = Inches(0.5)
    r2 = p_num2.add_run("2. Logika Reduksi Data (Reduction Logic): ")
    r2.bold = True
    p_num2.add_run(
        "Data transaksi kasir harian yang berjumlah ribuan baris merupakan data mentah berdimensi tinggi yang belum terstruktur untuk pengambilan keputusan. "
        "Sistem membutuhkan logika reduksi dan agregasi terprogram (Reduction Logic) untuk: (a) melakukan deduplikasi transaksi kasir dengan kompleksitas O(1) "
        "guna menyaring data redundan, (b) mereduksi riwayat transaksi multi-baris menjadi metrik ringkas per pelanggan (total belanja, akumulasi poin, dan frekuensi transaksi), "
        "serta (c) mereduksi seluruh dataset pelanggan menjadi daftar peringkat pemenang Top-N secara objektif melalui algoritma kampanye berbasis multi-kriteria "
        "dan mekanisme tie-breaker berjenjang (Total Belanja -> Frekuensi Belanja -> Waktu Pendaftaran)."
    )
    
    p = doc.add_paragraph(
        "Dengan mengintegrasikan arsitektur transaksi atomik dan logika reduksi data ke dalam aplikasi berbasis web, Toko Miniposh dapat mengeliminasi proses "
        "rekapitulasi manual dari kasir tunggal, mencegah kesalahan data, meningkatkan transparansi perolehan poin pelanggan, serta menetapkan penerima reward parsel "
        "secara cepat, adil, objektif, dan dapat dipertanggungjawabkan."
    )

    add_styled_heading(doc, "1.2 Rumusan Masalah (Naskah Hasil Revisi Lengkap)", level=2)
    
    p_rm_intro = doc.add_paragraph("Berdasarkan latar belakang dan identifikasi masalah di atas, maka rumusan masalah dalam penelitian ini dirumuskan sebagai berikut:")
    
    rms = [
        "Bagaimana merancang dan membangun sistem informasi loyalitas dan reward pelanggan berbasis web yang mampu mengintegrasikan data transaksi dari kasir tunggal (StoreTransactionSource) secara terstruktur pada Toko Miniposh?",
        "Bagaimana mengimplementasikan mekanisme transaksi atomik (Atomic Transactions) untuk menjamin integritas dan konsistensi data poin pada proses sinkronisasi transaksi massal dan penukaran reward?",
        "Bagaimana merancang dan menerapkan logika reduksi data (Reduction Logic) serta algoritma multi-level tie breaker untuk menyaring duplikasi transaksi dan menentukan pemenang reward campaign secara otomatis dan objektif?",
        "Bagaimana dampak implementasi sistem reward berbasis web terhadap penyelesaian gap loyalitas pelanggan dan efisiensi operasional evaluasi reward pada Toko Miniposh?"
    ]
    for i, rm in enumerate(rms, 1):
        p_rm = doc.add_paragraph()
        p_rm.paragraph_format.left_indent = Inches(0.4)
        p_rm.add_run(f"{i}. {rm}")
        
    doc.add_page_break()

    # =========================================================================
    # BAGIAN 2: LANDASAN KONSEP & TEORI (BAB II)
    # =========================================================================
    add_styled_heading(doc, "BAGIAN 2: LANDASAN KONSEP & TEORI BARU (BAB II)", level=1)
    
    add_callout(doc, "PETUNJUK PENEMPATAN DI FILE WORD SKRIPSI (BAB II)", 
                "Tambahkan sub-bab baru pada BAB II TINJAUAN PUSTAKA berikut untuk memperkuat jawaban atas pertanyaan dosen mengenai konsep 'Atomic' dan 'Reduction Logic' serta 'Gap Loyalitas Pelanggan'.\n"
                "Letakkan setelah Sub-bab 2.6 (Rule-Based System) atau sebelum Sub-bab 2.8.")
                
    add_styled_heading(doc, "2.6.1 Konsep Transaksi Atomik (Atomic Transactions)", level=2)
    p = doc.add_paragraph(
        "Transaksi atomik merupakan salah satu pilar utama dari prinsip ACID (Atomicity, Consistency, Isolation, Durability) dalam manajemen basis data modern. "
        "Sifat Atomicity menyatakan bahwa sebuah transaksi basis data yang terdiri atas sekumpulan operasi DML (Insert, Update, Delete) harus diperlakukan sebagai "
        "satu unit logika tunggal yang tidak dapat dipisahkan (indivisible). Konsep ini menganut prinsip all-or-nothing, di mana semua operasi dalam transaksi harus "
        "berhasil dieksekusi secara keseluruhan (commit), atau jika terjadi satu saja kegagalan, seluruh perubahan yang telah dilakukan akan dibatalkan secara total (rollback) "
        "(Silberschatz et al., 2020)."
    )
    p = doc.add_paragraph(
        "Dalam arsitektur sistem informasi reward Toko Miniposh, transaksi atomik diimplementasikan melalui modul Prisma ORM ($transaction). "
        "Penerapan transaksi atomik ini sangat penting pada dua proses kritis:"
    )
    p_at1 = doc.add_paragraph()
    p_at1.paragraph_format.left_indent = Inches(0.4)
    p_at1.add_run("a. Batch Synchronization & Point Accrual: ").bold = True
    p_at1.add_run("Proses pencatatan transaksi massal dari kasir yang memperbarui data Member (totalPoints, totalSpent, transactionCount) dan menyisipkan baris Transaction secara simultan.")
    
    p_at2 = doc.add_paragraph()
    p_at2.paragraph_format.left_indent = Inches(0.4)
    p_at2.add_run("b. Reward Redemption: ").bold = True
    p_at2.add_run("Proses penukaran reward yang memotong saldo poin member, memotong stok pada katalog reward, dan menerbitkan kode klaim penukaran secara terisolasi guna menghindari race condition.")

    add_styled_heading(doc, "2.6.2 Logika Reduksi Data (Reduction Logic) dalam Sistem Ritel", level=2)
    p = doc.add_paragraph(
        "Logika reduksi (Reduction Logic) merujuk pada prosedur algoritmik fungsional yang mentransformasikan sekumpulan data mentah berukuran besar (koleksi baris transaksi) "
        "menjadi struktur data berdimensi lebih ringkas yang merepresentasikan nilai akumulatif atau keputusan bisnis (Gamma et al., 1995; Fowler, 2018). "
        "Secara matematis, operasi reduksi memetakan himpunan transaksi T = {t_1, t_2, ..., t_n} milik seorang pelanggan menjadi vektor profil ringkas:"
    )
    
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_eq = p_eq.add_run("totalSpent = ∑ Amount(t_i),   totalPoints = ∑ Points(t_i),   transactionCount = |T|")
    r_eq.bold = True
    r_eq.italic = True
    
    p = doc.add_paragraph(
        "Dalam sistem ini, logika reduksi diterapkan pada tiga level:"
    )
    p_rl1 = doc.add_paragraph()
    p_rl1.paragraph_format.left_indent = Inches(0.4)
    p_rl1.add_run("1. Deduplication Reduction: ").bold = True
    p_rl1.add_run("Penyaringan data transaksi kasir yang redundan menggunakan hash set O(1) berbasis kombinasi (MemberId + Timestamp + Amount).")
    
    p_rl2 = doc.add_paragraph()
    p_rl2.paragraph_format.left_indent = Inches(0.4)
    p_rl2.add_run("2. Member Profile Aggregation: ").bold = True
    p_rl2.add_run("Penggabungan delta transaksi belanja baru ke dalam ringkasan statistik pelanggan di tabel Member.")
    
    p_rl3 = doc.add_paragraph()
    p_rl3.paragraph_format.left_indent = Inches(0.4)
    p_rl3.add_run("3. Campaign Winner Determination: ").bold = True
    p_rl3.add_run("Mereduksi ribuan baris transaksi dalam rentang periode kampanye menjadi daftar peringkat pemenang (Top-N) berdasarkan kriteria tertinggi dan aturan pemecah seri (tie-breaker).")

    add_styled_heading(doc, "2.7 Analisis Gap & Impact terhadap Loyalitas Pelanggan", level=2)
    
    table_gap = doc.add_table(rows=1, cols=3)
    table_gap.alignment = WD_TABLE_ALIGNMENT.CENTER
    ghdr = table_gap.rows[0].cells
    ghdr[0].text = "Dimensi Evaluasi"
    ghdr[1].text = "Kondisi Eksisting (Sebelum Sistem / Gap)"
    ghdr[2].text = "Kondisi Setelah Sistem (Impact Nyata)"
    
    for c in ghdr:
        set_cell_background(c, "1F4E79")
        set_cell_margins(c, 100, 100, 120, 120)
        p = c.paragraphs[0]
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9.5)
        
    gap_data = [
        ("Integrasi Transaksi", "Data belanja hanya tersimpan di log kasir tunggal (POS lokal); tidak terintegrasi ke data loyalitas.", "Tersedia modul Store Sync & Upload CSV yang menghubungkan data kasir langsung ke profil member."),
        ("Kecepatan Rekapitulasi", "Manual menggunakan spreadsheet atau ingatan; butuh waktu berhari-hari saat periode parsel Lebaran.", "Otomatis dan instan (dalam hitungan detik) mengevaluasi ribuan transaksi menjadi peringkat pemenang."),
        ("Objektivitas & Keadilan", "Rentan bias subjektif, salah catat, dan pilih kasih dalam menentukan penerima hadiah.", "100% objektif berbasis data riil transaksi, multi-kriteria transparan, dan algoritma tie-breaker teruji."),
        ("Transparansi Pelanggan", "Pelanggan tidak mengetahui saldo poin, status transaksi, maupun katalog reward yang tersedia.", "Tersedia portal member mandiri untuk cek saldo poin, riwayat transaksi, dan katalog reward."),
        ("Konsistensi Data", "Risiko ketidaksinkronan saldo poin saat terjadi pembatalan transaksi belanja.", "Dijamin melalui Atomic Transactions ($transaction) yang mencegah kegagalan parsial dan duplikasi data.")
    ]
    
    for row_idx, (d1, d2, d3) in enumerate(gap_data):
        row = table_gap.add_row().cells
        bg = "F2F7FA" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate([d1, d2, d3]):
            row[i].text = text
            set_cell_background(row[i], bg)
            set_cell_margins(row[i], 80, 80, 100, 100)
            p = row[i].paragraphs[0]
            p.runs[0].font.size = Pt(9)
            p.paragraph_format.line_spacing = 1.15
            
    doc.add_page_break()

    # =========================================================================
    # BAGIAN 3: PERBAIKAN DIAGRAM & RANCANGAN SISTEM (BAB III)
    # =========================================================================
    add_styled_heading(doc, "BAGIAN 3: PERBAIKAN DIAGRAM SISTEM (BAB III)", level=1)
    
    add_callout(doc, "PETUNJUK PENEMPATAN DI FILE WORD SKRIPSI (BAB III)", 
                "Buka file skripsi Anda pada 'BAB III METODOLOGI PENELITIAN'.\n"
                "1. Update Gambar 3.3 ERD Diagram dengan gambar ERD hasil import XML 'docs/erd_revisi_v2.xml'.\n"
                "2. Update Gambar 3.5 Class Diagram dengan gambar Class Diagram hasil import XML 'docs/class_diagram_revisi_v2.xml'.\n"
                "3. Update Gambar 3.6 Flowchart System dengan gambar Flowchart hasil import XML 'docs/flowchart_revisi_v2.xml'.\n"
                "4. Lengkapi narasi penjelasannya menggunakan teks perbaikan di bawah ini.")

    add_styled_heading(doc, "3.5.1 Perbaikan Entity Relationship Diagram (ERD)", level=2)
    p = doc.add_paragraph(
        "Entity Relationship Diagram (ERD) dirancang untuk memodelkan struktur data logis dan relasi fisik antar entitas basis data "
        "yang digunakan dalam sistem reward Toko Miniposh. Relasi utama yang dimodelkan meliputi:"
    )
    erd_points = [
        ("Member (1) ke Transaction (N): ", "Satu member dapat memiliki banyak riwayat transaksi belanja. Atribut foreign key memberId pada Transaction merujuk ke id pada entitas Member dengan aturan referensial Cascade Delete."),
        ("Member (1) ke RewardWinner (N): ", "Satu member berpotensi memenangkan banyak program reward campaign pada periode berbeda."),
        ("RewardCampaign (1) ke RewardWinner (N): ", "Satu program kampanye reward menghasilkan banyak baris pemenang sesuai konfigurasi winnersCount."),
        ("Member (1) ke RewardRedemption (N): ", "Satu member dapat melakukan penukaran reward berkali-kali selama saldo poin mencukupi."),
        ("RewardCatalog (1) ke RewardRedemption (N): ", "Satu item katalog reward dapat ditukarkan oleh banyak member (selama stok tersedia)."),
        ("StoreTransactionSource (Staging Table): ", "Berfungsi sebagai tabel perantara (staging table) yang menampung data mentah transaksi dari kasir toko eksternal (POS) sebelum melalui proses ETL (Extract, Transform, Load), deduplikasi, dan agregasi poin.")
    ]
    for b_title, b_desc in erd_points:
        p_b = doc.add_paragraph()
        p_b.paragraph_format.left_indent = Inches(0.4)
        r = p_b.add_run(b_title)
        r.bold = True
        p_b.add_run(b_desc)

    add_styled_heading(doc, "3.5.2 Perbaikan Class Diagram", level=2)
    p = doc.add_paragraph(
        "Class diagram menggambarkan struktur kelas perangkat lunak, atribut, method, serta hubungan asosiasi antar komponen sistem. "
        "Struktur kelas yang dibangun meliputi entitas model inti (Admin, Member, Transaction, SystemSetting, RewardCampaign, RewardWinner, "
        "RewardCatalog, RewardRedemption, dan StoreTransactionSource) serta kelas controller/service yang mengisolasi proses bisnis "
        "(TransactionService, RewardService, dan SyncService) dengan pemanfaatan Prisma ORM."
    )

    add_styled_heading(doc, "3.5.3 Perbaikan Flowchart Sistem", level=2)
    p = doc.add_paragraph(
        "Flowchart sistem menggambarkan alur kerja komprehensif mulai dari autentikasi admin, sinkronisasi data kasir (Store Sync / CSV Upload), "
        "validasi & deduplikasi O(1), kalkulasi poin atomik, konfigurasi kampanye reward multi-kriteria, penentuan pemenang dengan tie breaker, "
        "hingga manajemen klaim reward dan pencetakan laporan."
    )
    
    doc.add_page_break()

    # =========================================================================
    # BAGIAN 4: KODE XML DIAGRAM SIAP IMPORT (DRAW.IO)
    # =========================================================================
    add_styled_heading(doc, "BAGIAN 4: KODE XML DIAGRAM UNTUK DRAW.IO (DIAGRAMS.NET)", level=1)
    
    add_callout(doc, "CARA MENGGUNAKAN KODE XML DI DRAW.IO (DIAGRAMS.NET)", 
                "Langkah-langkah membuat gambar diagram rapi:\n"
                "1. Buka browser dan akses website https://app.diagrams.net/ (draw.io).\n"
                "2. Klik menu 'File' -> 'Open From' -> 'Device' (pilih file XML di folder docs/) ATAU klik 'Arrange' -> 'Insert' -> 'Advanced' -> 'XML' dan paste kode XML di bawah.\n"
                "3. Diagram akan otomatis tersusun dengan garis rapi, warna profesional, dan kardinalitas Crow's Foot.\n"
                "4. Export diagram ke format PNG/JPEG dengan resolusi tinggi (300 DPI) dan masukkan ke dokumen Skripsi Anda.")
                
    p = doc.add_paragraph(
        "File XML telah dibuatkan secara otomatis di dalam folder 'docs/' pada direktori proyek Anda:\n"
        "1. docs/erd_revisi_v2.xml (File ERD Lengkap & Rapi)\n"
        "2. docs/flowchart_revisi_v2.xml (File Flowchart Sistem Lengkap)\n"
        "3. docs/class_diagram_revisi_v2.xml (File Class Diagram Lengkap)\n"
    )
    
    doc.save("D:\\SKRIPSI\\reward_app_fiks\\DOKUMEN_REVISI_LENGKAP_SKRIPSI.docx")
    print("Successfully generated DOKUMEN_REVISI_LENGKAP_SKRIPSI.docx")

if __name__ == "__main__":
    create_full_revision_docx()
